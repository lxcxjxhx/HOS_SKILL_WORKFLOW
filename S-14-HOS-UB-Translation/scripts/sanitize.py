# -*- coding: utf-8 -*-
"""脱敏 docx 中的联系方式（邮箱、电话 -> __________），满足 Upwork 等平台"合同开始前政策"。
用法: python sanitize.py file.docx [更多.docx ...] [--no-phone]
默认规则：邮箱正则 + 9 位以上连续数字（视为电话；许可编号等 7-8 位数字不受影响）。
若需脱敏 7-8 位电话，先人工确认后用 --phone-min 7 调整。"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document

MASK = '__________'
EMAIL_RE = re.compile(r'[\w.+-]+@[\w-]+\.[\w.]+')


def main(files, phone_min=9):
    phone_re = re.compile(rf'(?<!\d)\d{{{phone_min},}}(?!\d)')

    def fix_run(r):
        if not r.text:
            return False
        if EMAIL_RE.search(r.text) or phone_re.search(r.text):
            t = EMAIL_RE.sub(MASK, r.text)
            t = phone_re.sub(MASK, t)
            r.text = t
            return True
        return False

    for f in files:
        d = Document(f)
        n = 0
        for p in d.paragraphs:
            for r in p.runs:
                n += fix_run(r)
        for tb in d.tables:
            for row in tb.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            n += fix_run(r)
        d.save(f)
        print('sanitized:', f, '| runs changed:', n)


if __name__ == '__main__':
    args = sys.argv[1:]
    phone_min = 9
    if '--no-phone' in args:
        args.remove('--no-phone')
        phone_min = 10 ** 9  # 相当于禁用电话规则
    if '--phone-min' in args:
        i = args.index('--phone-min')
        phone_min = int(args[i + 1])
        del args[i:i + 2]
    if not args:
        print(__doc__)
        sys.exit(1)
    main(args, phone_min=phone_min)
