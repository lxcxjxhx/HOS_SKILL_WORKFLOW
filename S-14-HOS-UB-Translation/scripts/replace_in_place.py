# -*- coding: utf-8 -*-
"""原地替换 docx 文本为译文，保留段落/字体/表格/分页格式。
用法: python replace_in_place.py input.docx output.docx translation_map.py
translation_map.py 内容示例:
    TOP = {0: "标题译文", 15: "含\\n段内换行的译文"}
    TBL = {(0, 0, 0): "表格译文"}
（TOP 键=顶层段索引；TBL 键=(行,列,单元格内段序)。未列出的段落保持原文。
 排版微调：正文右对齐(RTL 残留)改为两端对齐，标题保持居中——唯一允许的调整。）"""
import sys, copy, importlib.util
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

CONTENT_TAGS = (qn('w:r'), qn('w:br'), qn('w:hyperlink'), qn('w:ins'), qn('w:del'),
                qn('w:drawing'), qn('w:pict'), qn('w:object'), qn('w:smartTag'))


def set_text(p, text):
    """清空段落内容（保留 pPr），按首个 run 的 rPr 重建单 run，写入译文（\\n -> w:br）"""
    rpr = None
    for r in p.runs:
        if r._element.rPr is not None:
            rpr = copy.deepcopy(r._element.rPr)
            break
    for child in list(p._element):
        if child.tag in CONTENT_TAGS:
            p._element.remove(child)
    if text == '':
        return
    run = p.add_run()
    if rpr is not None:
        run._element.insert(0, rpr)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    if not rFonts.get(qn('w:eastAsia')):
        rFonts.set(qn('w:eastAsia'), '宋体')
    for i, part in enumerate(text.split('\n')):
        if i > 0:
            run._element.append(OxmlElement('w:br'))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = part
        run._element.append(t)


def main(src, dst, map_path):
    spec = importlib.util.spec_from_file_location('tmap', map_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    TOP = getattr(mod, 'TOP', {})
    TBL = getattr(mod, 'TBL', {})
    d = Document(src)
    for i, p in enumerate(d.paragraphs):
        if i in TOP:
            set_text(p, TOP[i])
        if i in TOP and p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for tb in d.tables:
        for ri, row in enumerate(tb.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    key = (ri, ci, pi)
                    if key in TBL:
                        set_text(p, TBL[key])
                    if key in TBL and p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    d.save(dst)
    print('saved:', dst)


if __name__ == '__main__':
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3])
