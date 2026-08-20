# -*- coding: utf-8 -*-
"""提取 docx 全部段落与表格文本（标注段内换行 [BR]），供翻译规划。
用法: python extract_docx.py input.docx"""
import sys, io
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document


def main(path):
    d = Document(path)
    print(f'== 顶层段落（{len(d.paragraphs)}）==')
    for i, p in enumerate(d.paragraphs):
        t = p.text.replace('\n', ' [BR] ')
        print(f'{i:3d}|{t}')
    for ti, tb in enumerate(d.tables):
        print(f'== 表格 {ti}（{len(tb.rows)}x{len(tb.columns)}）==')
        for ri, row in enumerate(tb.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    if p.text.strip():
                        print(f'r{ri}c{ci}p{pi}|{p.text.replace(chr(10), " [BR] ")}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    main(sys.argv[1])
